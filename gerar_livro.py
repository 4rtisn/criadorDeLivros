import time
import os
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import pyperclip
from docx import Document
from bs4 import BeautifulSoup
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm


def verificar_pausa_gemini(driver):
    try:
        mensagem_pausa = driver.find_element(By.XPATH, "//h1[contains(text(), 'O Gemini está dando uma pausa.')]")
        return True if mensagem_pausa else False
    except:
        return False

def lidar_com_pausa_gemini(driver):
    print("Pausa detectada. Aguardando 30 minutos antes de tentar novamente...")
    time.sleep(1800)  
    driver.refresh()
    time.sleep(20)  


def clicar_ultimo_botao(driver, seletor):
    botoes = driver.find_elements(By.CSS_SELECTOR, seletor)
    if botoes:
        ultimo_botao = botoes[-1]
        ultimo_botao.click()
    else:
        raise Exception(f"Botão '{seletor}' não encontrado.")

def remover_titulos(texto):
    linhas = texto.split("\n")
    linhas_sem_titulos = [linha for linha in linhas if not linha.startswith("### **")]
    return "\n".join(linhas_sem_titulos)

def gerar_escopo(driver, tema):
    try:
        driver.refresh() 
        time.sleep(2)
        
        text_area = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.CSS_SELECTOR, "[data-placeholder='Digite uma pergunta ou comando']"))
        )
        text_area.clear()
        pergunta = f"30 dicas sobre {tema}"
        text_area.send_keys(pergunta)
        text_area.send_keys('\n')

        time.sleep(30)

        clicar_ultimo_botao(driver, "[mattooltip='Mais']")
        time.sleep(2)
        clicar_ultimo_botao(driver, "[aria-label='Copiar']")
        time.sleep(2)

        texto_da_area_de_transferencia = pyperclip.paste()
        
        passos = [linha.strip() for linha in texto_da_area_de_transferencia.split("\n") if linha.strip() and not linha.startswith("###")]
        
        return passos

    except Exception as e:
        print(f"Erro ao gerar o escopo para o tema '{tema}': {e}")
        return []

def gerar_dissertacao(driver, tema, topico):
    try:
        driver.refresh() 
        time.sleep(2)
        
        text_area = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.CSS_SELECTOR, "[data-placeholder='Digite uma pergunta ou comando']"))
        )
        text_area.clear()
        pergunta = f"Para o livro '{tema}', disserte sobre {topico}. Narrativa completa."
        text_area.send_keys(pergunta)
        text_area.send_keys('\n')

        time.sleep(60)

        clicar_ultimo_botao(driver, "[mattooltip='Mais']")
        time.sleep(2)
        clicar_ultimo_botao(driver, "[aria-label='Copiar']")
        time.sleep(2)

        texto_da_area_de_transferencia = pyperclip.paste()

        return remover_titulos(texto_da_area_de_transferencia)

    except Exception as e:
        print(f"Erro ao gerar dissertação para o tópico '{topico}': {e}")
        return ""

def formatar_documento(doc):
    section = doc.sections[0]
    section.page_height = Mm(210)
    section.page_width = Mm(148)
    section.left_margin = Mm(6.4)
    section.right_margin = Mm(6.4)
    section.top_margin = Mm(6.4)
    section.bottom_margin = Mm(6.4)
    section.gutter = Mm(12.7)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(14)

def salvar_em_docx_imediato(titulo, conteudo, doc, livro_nome):
    doc.add_paragraph(titulo, style='Heading1').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    soup = BeautifulSoup(conteudo, 'html.parser')
    for element in soup.stripped_strings:
        doc.add_paragraph(element, style='Normal')

    doc.save(livro_nome)

def tema_ja_gerado(tema, arquivo_gerados):
    if not os.path.exists(arquivo_gerados):
        return False
    
    with open(arquivo_gerados, 'r') as f:
        temas_gerados = f.read().splitlines()
    
    return tema in temas_gerados

def registrar_tema_gerado(tema, arquivo_gerados):
    with open(arquivo_gerados, 'a') as f:
        f.write(tema + '\n')

def obter_proximo_tema(arquivo_temas, arquivo_gerados):
    with open(arquivo_temas, 'r') as f:
        temas = f.read().splitlines()

    for tema in temas:
        if not tema_ja_gerado(tema, arquivo_gerados):
            return tema
    
    return None

def criar_livro():
    chrome_options = Options()
    chrome_options.add_argument("--disable-blink-features=AutomationControlled")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--disable-gpu")

    driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()), options=chrome_options)
    driver.get("https://bard.google.com/")

    input("Faça login e resolva o captcha manualmente. Pressione Enter quando estiver pronto...")

    try:
        arquivo_temas = "temas.txt"
        arquivo_gerados = "temas_gerados.txt"

        tema = obter_proximo_tema(arquivo_temas, arquivo_gerados)
        if not tema:
            print("Todos os temas disponíveis já foram gerados.")
            return
        
        doc = Document() 
        formatar_documento(doc) 

        doc.add_paragraph(tema, style='Title').alignment = WD_ALIGN_PARAGRAPH.CENTER  

        livro_nome = f"livro_{tema.replace(' ', '_')}.docx"

        escopo = gerar_escopo(driver, tema)
        if not escopo:
            print(f"Não foi possível gerar o escopo para o tema '{tema}'.")
            return

        for i, topico in enumerate(escopo):
            if verificar_pausa_gemini(driver):
                lidar_com_pausa_gemini(driver)

            print(f"Gerando dissertação para o tópico {i + 1}: {topico}")
            dissertacao = gerar_dissertacao(driver, tema, topico)
            if dissertacao:
                salvar_em_docx_imediato(f"Capítulo {i + 1}: {topico}", dissertacao, doc, livro_nome)

        registrar_tema_gerado(tema, arquivo_gerados)

        print(f"Livro '{livro_nome}' criado com sucesso.")

    except Exception as e:
        print(f"Erro durante o processo: {e}")

    finally:
        driver.quit()

if __name__ == "__main__":
    criar_livro()
